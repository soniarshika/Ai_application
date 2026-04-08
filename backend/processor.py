"""
Document processor for logistics documents.

Chunking strategy — three content types, each handled differently:

1. KV BLOCK  (e.g. "Shipper: ACME Corp\nAddress: 123 Main St\nCity: Chicago")
   ─ Consecutive Key: Value lines (≥2) are kept as ONE atomic chunk.
   ─ Why: A question like "Who is the shipper?" needs the full shipper block
     (name + address + city) in a single retrieval hit. Splitting them means
     cosine similarity finds the name chunk but misses the address.

2. TABLE ROW (e.g. "[TABLE] Weight: 4200 lbs | Class: 70 | Commodity: Steel")
   ─ Each table row is one chunk, headers prepended as "Header: value" pairs.
   ─ Why: Table rows are already structured records. They embed well as-is and
     align with how users phrase queries ("what is the weight of the shipment?").

3. NARRATIVE  (clauses, terms, conditions, remarks)
   ─ Sentence-boundary split into 500-char windows with 100-char overlap.
   ─ Why: Free-form text needs overlap so a sentence straddling a boundary
     isn't split mid-thought. 500 chars ≈ 80–100 tokens, the sweet spot for
     text-embedding-3-small — specific enough to score high cosine similarity,
     large enough to carry full semantic context.

Key design decisions:
  - MIN_CHUNK = 15 chars (not 300) so short but critical KV lines are NOT dropped.
  - KV lines are grouped, not stored individually.
  - Table rows get a "TABLE:" prefix for retrieval clarity.
  - No importance scoring or keyword extraction — this adds noise, not signal,
    to the embedding. Embeddings encode semantics; metadata scoring is redundant.
"""

import re
import io
import logging
from typing import List, Tuple, Dict

log = logging.getLogger("logistics.processor")

# Narrative chunks: 500 chars with 100-char overlap
CHUNK_SIZE    = 500
CHUNK_OVERLAP = 100
MIN_CHUNK     = 15   # minimum chars — keeps short but critical KV lines

# Matches logistics key-value lines: "Shipper Name: ACME Corp", "Rate: $1,250", etc.
KV_LINE_RE = re.compile(r"^[\w\s\-\.\/()\[\]]{1,50}:\s+\S.*$")


class DocumentProcessor:

    def parse(self, file_bytes: bytes, filename: str) -> Tuple[str, List[Dict]]:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        log.debug(f"  Parsing {ext.upper()} ({len(file_bytes):,} bytes)")
        if ext == "pdf":
            return self._parse_pdf(file_bytes)
        elif ext == "docx":
            return self._parse_docx(file_bytes)
        else:
            return self._parse_txt(file_bytes)

    # ── Parsers ──────────────────────────────────────────────────────────

    def _parse_pdf(self, data: bytes) -> Tuple[str, List[Dict]]:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=data, filetype="pdf")
        page_texts = []

        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text") or ""

            # Extract tables and append as structured [TABLE] lines
            try:
                for table in page.find_tables():
                    rows = table.extract()
                    if not rows:
                        continue
                    headers = [str(h).strip() if h else f"Col{i}" for i, h in enumerate(rows[0])]
                    for row in rows[1:]:
                        cells = [
                            f"{headers[i]}: {str(cell).strip()}"
                            for i, cell in enumerate(row)
                            if cell and str(cell).strip()
                        ]
                        if cells:
                            text += "\n[TABLE] " + " | ".join(cells)
            except Exception:
                pass

            page_texts.append({"page": page_num, "text": text})

        doc.close()
        full_text = "\n".join(p["text"] for p in page_texts)
        return full_text, page_texts

    def _parse_docx(self, data: bytes) -> Tuple[str, List[Dict]]:
        from docx import Document

        doc = Document(io.BytesIO(data))
        lines = []

        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)

        for table in doc.tables:
            if not table.rows:
                continue
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            for row in table.rows[1:]:
                cells = [
                    f"{headers[i]}: {cell.text.strip()}"
                    for i, cell in enumerate(row.cells)
                    if cell.text.strip()
                ]
                if cells:
                    lines.append("[TABLE] " + " | ".join(cells))

        full_text = "\n".join(lines)
        return full_text, [{"page": 1, "text": full_text}]

    def _parse_txt(self, data: bytes) -> Tuple[str, List[Dict]]:
        text = data.decode("utf-8", errors="replace")
        return text, [{"page": 1, "text": text}]

    # ── Chunking ─────────────────────────────────────────────────────────

    def chunk_text(self, page_texts: List[Dict]) -> List[Dict]:
        """
        Produce a flat list of chunk dicts:
          { text, page_number, chunk_type, chunk_index }
        """
        all_chunks: List[Dict] = []

        for page_info in page_texts:
            page_num = page_info["page"]
            lines = page_info["text"].split("\n")
            log.debug(f"  Chunking page {page_num} — {len(lines)} lines")
            chunks = self._process_lines(lines, page_num)
            log.debug(f"    → {len(chunks)} chunks from page {page_num}")
            all_chunks.extend(chunks)

        # Sequential reindex
        for i, chunk in enumerate(all_chunks):
            chunk["chunk_index"] = i

        from collections import Counter
        breakdown = Counter(c["chunk_type"] for c in all_chunks)
        log.debug(
            f"  Chunk breakdown: total={len(all_chunks)}  " +
            "  ".join(f"{t}={n}" for t, n in sorted(breakdown.items()))
        )
        return all_chunks

    def _process_lines(self, lines: List[str], page_num: int) -> List[Dict]:
        chunks: List[Dict] = []
        kv_buffer: List[str] = []
        narrative_buffer: List[str] = []

        def flush_kv():
            if kv_buffer:
                text = "\n".join(kv_buffer)
                if len(text) >= MIN_CHUNK:
                    chunks.append(_chunk(text, page_num, "kv_block"))
                kv_buffer.clear()

        def flush_narrative():
            if narrative_buffer:
                text = " ".join(narrative_buffer).strip()
                narrative_buffer.clear()
                if not text or len(text) < MIN_CHUNK:
                    return
                for part in _split_narrative(text):
                    chunks.append(_chunk(part, page_num, "narrative"))

        for raw_line in lines:
            line = raw_line.strip()

            if not line:
                # Blank line → paragraph boundary
                flush_kv()
                flush_narrative()
                continue

            if line.startswith("[TABLE]"):
                flush_kv()
                flush_narrative()
                content = "TABLE: " + line[7:].strip()
                if len(content) >= MIN_CHUNK:
                    chunks.append(_chunk(content, page_num, "table_row"))
                continue

            if KV_LINE_RE.match(line):
                flush_narrative()
                kv_buffer.append(line)
                continue

            # If we have a pending KV buffer and this line doesn't match KV,
            # flush and start narrative
            flush_kv()
            narrative_buffer.append(line)

        flush_kv()
        flush_narrative()
        return chunks


# ── Module-level helpers (pure functions) ────────────────────────────────

def _chunk(text: str, page_num: int, chunk_type: str) -> Dict:
    return {
        "text":        text,
        "page_number": page_num,
        "chunk_type":  chunk_type,
        "chunk_index": 0,  # overwritten by caller
    }


def _split_narrative(text: str) -> List[str]:
    """
    Sentence-boundary split with sliding window overlap.
    Window: CHUNK_SIZE chars, overlap: CHUNK_OVERLAP chars.
    """
    # Split on sentence boundaries
    sentences = re.split(r'(?<=[.!?])\s+', text)

    chunks: List[str] = []
    current = ""

    for sent in sentences:
        if not sent.strip():
            continue
        candidate = (current + " " + sent).strip() if current else sent
        if len(candidate) <= CHUNK_SIZE:
            current = candidate
        else:
            if current and len(current) >= MIN_CHUNK:
                chunks.append(current)
            # Start new chunk with overlap from end of previous
            if current and len(current) > CHUNK_OVERLAP:
                overlap_text = current[-CHUNK_OVERLAP:]
                current = (overlap_text + " " + sent).strip()
            else:
                current = sent

    if current and len(current) >= MIN_CHUNK:
        chunks.append(current)

    # If text was too short to split, return as-is
    if not chunks and len(text) >= MIN_CHUNK:
        chunks.append(text)

    return chunks
